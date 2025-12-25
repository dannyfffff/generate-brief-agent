from typing import Dict, Any, Optional, Literal, List
from state import WorkflowState
from tools import llm, search_tool
import os
import datetime
import re
from concurrent.futures import ThreadPoolExecutor, as_completed
from langchain_community.document_loaders import WebBaseLoader
from langchain_core.prompts import ChatPromptTemplate
from prompts import (BRIEFING_SYSTEM_PROMPT,
                     BRIEFING_HUMAN_PROMPT_TEMPLATE_A,
                     BRIEFING_HUMAN_PROMPT_TEMPLATE_B, BRIEFING_HUMAN_PROMPT_TEMPLATE_C,

                     CHAPTER_3_SYSTEM_PROMPT_TEMPLATE_A,
                     CHAPTER_3_SYSTEM_PROMPT_TEMPLATE_B, CHAPTER_3_SYSTEM_PROMPT_TEMPLATE_C,
                     CHAPTER_3_USER_PROMPT_TEMPLATE_A,
                     CHAPTER_3_USER_PROMPT_TEMPLATE_B, CHAPTER_3_USER_PROMPT_TEMPLATE_C,CHAPTER_2_SYSTEM_PROMPT_A_OUTLINE,
                     CHAPTER_2_USER_PROMPT_TEMPLATE_A_OUTLINE, CHAPTER_2_SYSTEM_PROMPT_B_OUTLINE,
                     CHAPTER_2_USER_PROMPT_TEMPLATE_B_OUTLINE, CHAPTER_2_SYSTEM_PROMPT_C_OUTLINE,
                     CHAPTER_2_USER_PROMPT_TEMPLATE_C_OUTLINE, CHAPTER_2_SYSTEM_PROMPT_A_AUTO,
                     CHAPTER_2_USER_PROMPT_TEMPLATE_A_AUTO, CHAPTER_2_SYSTEM_PROMPT_B_AUTO,
                     CHAPTER_2_USER_PROMPT_TEMPLATE_B_AUTO, CHAPTER_2_SYSTEM_PROMPT_C_AUTO,
                     CHAPTER_2_USER_PROMPT_TEMPLATE_C_AUTO
                     )
from docx import Document




# 核心决策节点：根据用户明确选择的 ID 填充 OptionContent
def question_node(state: WorkflowState) -> WorkflowState:
    """接收 optionId，并自动填充 optionContent"""

    # 从状态中获取规定的变量名 optionId
    selected_id = state['optionId']

    # 预设选项内容
    options_map = {
        "A": "政策类",
        "B": "技术类",
        "C": "事件类",
        "Other": "用户不可见/自由输入",
    }

    selected_content = options_map.get(selected_id, "未知选项")

    print(f"--- 问答节点 (question_node)：处理选择 {selected_id}，内容为 {selected_content} ---")

    # 更新状态，自动填充 optionContent
    return {
        "optionContent": selected_content,
    }


def _execute_site_search(query: str) -> list:
    """执行 Tavily 搜索并返回映射后的文档列表。"""
    try:
        tavily_results = search_tool.invoke({"query": query})

        doc_results_list = []
        for item in tavily_results:
            # 统一结果映射逻辑
            url = item.get('source') if hasattr(item, 'metadata') else item.get('url')
            doc_results_list.append({
                "title": item.get('title', '无标题') if hasattr(item, 'metadata') else item.get('title', '无标题'),
                "url": url,
                "content": item.page_content if hasattr(item, 'metadata') else item.get('content', '无内容'),
            })
        return doc_results_list
    except Exception as e:
        print(f"搜索失败: {e}")
        return []

# 站点配置列表
SITE_CONFIGS = [
    {"domain": "ithome.com", "prefix": "it", "modifier": "AI热点新闻"},
    {"domain": "qbitai.com", "prefix": "lz", "modifier": "AI技术热点"},
    {"domain": "aiera.com.cn", "prefix": "xz", "modifier": "AI政策"},
    {"domain": "jiqizhixin.com", "prefix": "jq", "modifier": "AI政策"},
]


def route_option_node(state: WorkflowState):
    return state  # 仅作路由，不修改状态

def Agent_Agent_node(state: WorkflowState) -> WorkflowState:
    """
    进入 A 节点，在函数内部顺序执行所有四个站点的独立搜索。
    结果分别存储到 it_doc_results, lz_doc_results, xz_doc_results, jq_doc_results。
    """
    workflow_input = state['workflowInput']

    # 初始化一个字典来收集所有搜索结果
    all_results = {}

    for config in SITE_CONFIGS:
        domain = config["domain"]
        prefix = config["prefix"]
        modifier = config["modifier"]
        output_key = f"{prefix}_doc_results"  # 构建输出状态键

        # 构造精确的搜索查询
        search_query = f'site:{domain} "{workflow_input}" {modifier}'

        print(f"*** 正在执行 {domain} 搜索。查询: {search_query} ***")

        # 1. 执行搜索
        doc_results_list = _execute_site_search(search_query)

        # 2. 存储结果
        all_results[output_key] = doc_results_list
        print(f"*** {domain} 搜索成功，找到 {len(doc_results_list)} 条结果。***")

    # 一次性将所有四个搜索结果更新到状态中
    return all_results



MEDIA_CONFIG = {
    "it": {"domain": "ithome.com", "input_key": "it_doc_results"},
    "lz": {"domain": "qbitai.com", "input_key": "lz_doc_results"},
    "xz": {"domain": "aiera.com.cn", "input_key": "xz_doc_results"},
    "jq": {"domain": "jiqizhixin.com", "input_key": "jq_doc_results"},
}


def create_fetch_all_data_node(config: Dict[str, Any]):
    """
    创建一个封装了所有媒体源的 URL 提取（12个）和内容抓取（12个）逻辑的单一节点。

    config: 使用 MEDIA_CONFIG 字典进行配置。
    """

    def fetch_all_data_node(state: WorkflowState) -> Dict[str, Any]:
        all_results = {}
        fetch_jobs = {}  # 用于存储并行任务

        # 使用多线程执行器加速网页抓取
        # 线程数设置为 12 (4个媒体 * 3个链接)，但限制在合理范围，如 max_workers=8
        with ThreadPoolExecutor(max_workers=8) as executor:

            for prefix, media_conf in config.items():
                input_key = media_conf["input_key"]
                domain = media_conf["domain"]
                doc_results = state.get(input_key, [])

                # 1. 内部执行 URL 提取和过滤 (原 12 个 xx_url_i 节点功能)
                valid_docs = [
                    doc for doc in doc_results
                    if domain in doc.get('url', '') and doc.get('content', '').strip()
                ]

                # 2. 针对前 3 个有效结果，启动并行内容抓取 (原 12 个 xx_fetch_i 节点功能)
                for i in range(1, 4):
                    content_key = f"{prefix}_content_{i}"
                    title_key = f"{prefix}_title_{i}"

                    if len(valid_docs) >= i:
                        target_doc = valid_docs[i - 1]
                        target_url = target_doc.get('url', 'URL获取失败')

                        # 提交抓取任务到线程池
                        future = executor.submit(
                            _fetch_single_content,
                            target_url,
                            content_key,
                            title_key
                        )
                        fetch_jobs[future] = (content_key, title_key)
                    else:
                        # 标记缺失结果
                        all_results[content_key] = "有效结果不足"
                        all_results[title_key] = "有效结果不足"

            # 3. 收集所有并行任务的结果
            for future in as_completed(fetch_jobs):
                try:
                    result = future.result()
                    all_results.update(result)
                except Exception as e:
                    print(f"并行抓取任务出现未预期错误: {e}")

        print(f"✅ Fetch_All_Data_Node: 成功封装并完成了所有 12 个 URL 的内容抓取。")
        # 返回所有 36 个状态键（content_i 和 title_i）
        return all_results

    return fetch_all_data_node


# --- 辅助函数：单个 URL 抓取逻辑 ---
def _fetch_single_content(url: str, content_key: str, title_key: str) -> Dict[str, Any]:
    """执行单个 URL 的抓取任务。"""
    if not url or not url.startswith("http"):
        return {content_key: "无有效 URL", title_key: "无有效 URL"}

    try:
        # ⚠️ 这里假设 WebBaseLoader 和 documents[0].page_content 逻辑是正确的
        loader = WebBaseLoader(url)
        documents = loader.load()
        full_content = documents[0].page_content if documents else "内容抓取失败"
        full_title = documents[0].metadata.get('title', '无标题') if documents else "无标题"

        return {
            content_key: full_content,
            title_key: full_title
        }
    except Exception as e:
        return {content_key: f"抓取失败: {e}", title_key: "抓取失败"}


def aggregate_and_draft_node(state: WorkflowState) -> WorkflowState:
    """
    合并 Step 1 和 Step 2，并调用 prompts.py 中的提示词模板。
    """
    workflow_input = state['workflowInput']
    option_id = state.get('optionId')
    # --- A. 内容聚合 (事实基础构建) ---
    aggregated_content_parts = []

    # ... (内容聚合逻辑保持不变，生成 raw_context) ...
    for i in range(1, 4):
        for prefix in ["it", "lz", "xz", "jq"]:
            content_key = f"{prefix}_content_{i}"
            title_key = f"{prefix}_title_{i}"

            content = state.get(content_key)
            title = state.get(title_key)

            if content and len(content) > 200 and "抓取失败" not in content:
                snippet = content[:3000]
                aggregated_content_parts.append(f"--- 文章标题: {title} ---\n内容摘要：{snippet}...\n")

    if not aggregated_content_parts:
        return {"briefing_draft": "未找到任何有效且抓取成功的文章内容，无法撰写简报。"}

    raw_context = "\n\n".join(aggregated_content_parts)

    human_template = BRIEFING_HUMAN_PROMPT_TEMPLATE_A

    if option_id == "A":
        human_template = BRIEFING_HUMAN_PROMPT_TEMPLATE_A
    elif option_id == "B":
        human_template = BRIEFING_HUMAN_PROMPT_TEMPLATE_B
    elif option_id == "C":
        human_template = BRIEFING_HUMAN_PROMPT_TEMPLATE_C

    human_prompt_content = human_template.format(
        raw_context=raw_context,
        workflow_input=workflow_input
    )

    final_prompt = ChatPromptTemplate.from_messages([
        ("system", BRIEFING_SYSTEM_PROMPT),
        ("human", human_prompt_content)
    ])

    print(f"--- 合并节点：聚合了 {len(aggregated_content_parts)} 篇文章，正在调用 LLM 一步撰写简报... ---")

    try:
        chain = final_prompt | llm
        result_message = chain.invoke({})

        return {"briefing_draft": result_message.content}
    except Exception as e:
        return {"error": f"内容聚合与简报撰写失败: {e}"}



def user_query_node(state: WorkflowState) -> WorkflowState:

    briefing = state.get('briefing_draft')

    # --- 1. 打印提示信息 (保留，用于日志追踪) ---

    query_text = (
        "如果给我大纲，能让我更准确地理解文章的重要侧重点。\n"
        "请问您有什么建议吗？"
    )

    full_prompt_to_user = (
        f"【已生成简报内容】\n\n{briefing}\n"
        f"\n\n【下一步操作建议】\n{query_text}"
    )

    print("\n" + "=" * 60)
    print("✅ [临时] 自动选择：A 选项")
    print(full_prompt_to_user)
    print("=" * 60 + "\n")

    return {
        "briefing_draft": briefing
    }


def generate_chapter_2_node(state: WorkflowState) -> WorkflowState:
    input1_content = state.get('briefing_draft', "")
    option_id = state.get('optionId')
    user_choice_id = state.get('userChoiceId')
    user_outline = state.get('userOutline', "")

    if not input1_content:
        return {"chapter_2_content": "错误：缺少事实基础，无法生成第二章。"}

    system_template = None
    user_template = None

    # 1. 外部判断：模式 (OUTLINE vs AUTO)
    # 只有 userChoiceId='B' 且大纲非空时才进入 OUTLINE 模式
    if user_choice_id == "B" and user_outline.strip():
        mode_suffix = "OUTLINE"

        # 2. 内部判断：领域 (A/B/C)
        if option_id == "A":
            system_template = CHAPTER_2_SYSTEM_PROMPT_A_OUTLINE
            user_template = CHAPTER_2_USER_PROMPT_TEMPLATE_A_OUTLINE
        elif option_id == "B":
            system_template = CHAPTER_2_SYSTEM_PROMPT_B_OUTLINE
            user_template = CHAPTER_2_USER_PROMPT_TEMPLATE_B_OUTLINE
        elif option_id == "C":
            system_template = CHAPTER_2_SYSTEM_PROMPT_C_OUTLINE
            user_template = CHAPTER_2_USER_PROMPT_TEMPLATE_C_OUTLINE

    # 其他所有情况 (包括 userChoiceId='A', 'Other', 或大纲为空) 都视为 AUTO 模式
    else:
        mode_suffix = "AUTO"

        # 2. 内部判断：领域 (A/B/C)
        if option_id == "A":
            system_template = CHAPTER_2_SYSTEM_PROMPT_A_AUTO
            user_template = CHAPTER_2_USER_PROMPT_TEMPLATE_A_AUTO
        elif option_id == "B":
            system_template = CHAPTER_2_SYSTEM_PROMPT_B_AUTO
            user_template = CHAPTER_2_USER_PROMPT_TEMPLATE_B_AUTO
        elif option_id == "C":
            system_template = CHAPTER_2_SYSTEM_PROMPT_C_AUTO
            user_template = CHAPTER_2_USER_PROMPT_TEMPLATE_C_AUTO

    # 3. 兜底检查
    if not system_template or not user_template:
        print(f"🚨 无法匹配提示词 ({option_id}_{mode_suffix})。使用默认 A_AUTO。")
        system_template = CHAPTER_2_SYSTEM_PROMPT_A_AUTO
        user_template = CHAPTER_2_USER_PROMPT_TEMPLATE_A_AUTO
        mode_suffix = "DEFAULT_AUTO"

    # 4. 构造 Prompt
    final_prompt = ChatPromptTemplate.from_messages([
        ("system", system_template),
        ("human", user_template)
    ])

    # 5. 构造变量字典：包含 input1 和大纲 (dagang 对应 userOutline)
    prompt_variables = {
        "input1": input1_content,
        "dagang": user_outline,
    }

    try:
        chain = final_prompt | llm
        result_message = chain.invoke(prompt_variables)

        generated_chapter_2 = result_message.content

        print(f"✅ 已使用 {option_id} 领域提示词，模式 {mode_suffix} 生成第二章内容。")
        return {"chapter_2_content": generated_chapter_2}

    except Exception as e:
        print(f"🚨 生成第二章失败: {e}")
        return {"chapter_2_content": f"错误：生成第二章失败: {e}"}



def generate_chapter_3_node(state: WorkflowState) -> WorkflowState:
    """
    基于事实基础、第二章内容和 lz_content_1 的精准资料，生成简报第三章 (对策建议)。
    """
    # 提取所需的所有信息
    input1 = state.get('briefing_draft', "")
    zhengce = state.get('lz_content_1', "")
    question = state.get('chapter_2_content', "")
    option_id = state.get('optionId')

    if not input1 or not zhengce or not question:
        # 添加调试信息，帮助判断哪个字段缺失
        error_msg = f"错误：生成第三章所需的基础信息缺失。简述:{len(input1) > 0}, 量子位:{len(zhengce) > 0}, 第二章:{len(question) > 0}"
        return {"chapter_3_content": error_msg}


    system_template = CHAPTER_3_SYSTEM_PROMPT_TEMPLATE_A  # 兜底
    human_template = CHAPTER_3_USER_PROMPT_TEMPLATE_A  # 兜底

    if option_id == "A":
        system_template = CHAPTER_3_SYSTEM_PROMPT_TEMPLATE_A
        human_template = CHAPTER_3_USER_PROMPT_TEMPLATE_A
    elif option_id == "B":
        system_template = CHAPTER_3_SYSTEM_PROMPT_TEMPLATE_B
        human_template = CHAPTER_3_USER_PROMPT_TEMPLATE_B
    elif option_id == "C":
        system_template = CHAPTER_3_SYSTEM_PROMPT_TEMPLATE_C
        human_template = CHAPTER_3_USER_PROMPT_TEMPLATE_C

    final_prompt = ChatPromptTemplate.from_messages([
        ("system", system_template),  # 包含 {input1}, {zhengce}, {question}
        ("human", human_template)  # 包含 {input1}, {zhengce}, {question}
    ])

    # 将所有变量打包成一个字典
    prompt_variables = {
        "input1": input1,
        "zhengce": zhengce,
        "question": question
    }

    # 2. 调用 LLM：将变量字典传递给 invoke()
    try:
        chain = final_prompt | llm
        result_message = chain.invoke(prompt_variables)  # <-- 传递变量字典

        generated_chapter_3 = result_message.content

        print("✅ 已生成第三章内容。")
        return {"chapter_3_content": generated_chapter_3}

    except Exception as e:
        print(f"🚨 生成第三章失败: {e}")
        return {"chapter_3_content": f"错误：生成第三章失败: {e}"}


# --- 3. 综合节点 ---
def combine_briefing_node(state: WorkflowState) -> WorkflowState:
    """
    严格按照用户要求，将简报的首段/事件简述、第二章和第三章内容按顺序拼接。
    """
    # 提取三段内容
    jianshu = state.get('briefing_draft', "")
    di2zhang = state.get('chapter_2_content', "")
    di3zhang = state.get('chapter_3_content', "")

    if not jianshu or not di2zhang or not di3zhang:
        return {"final_briefing": "错误：简报三部分内容不完整，无法拼接。"}

    # 严格按照用户要求：按顺序拼接，不修改任何文字、符号、段落
    final_output = jianshu + "\n\n" + di2zhang + "\n\n" + di3zhang

    print("✅ 已完成三章简报的最终拼接。")
    return {"final_briefing": final_output}



def export_to_docx_node(state: WorkflowState) -> WorkflowState:
    """
    接收 final_briefing 内容，生成 DOCX 文件，并将文件路径/内容存储在状态中。
    优化点：1. 使用标题+时间戳命名； 2. 统一保存到历史目录。
    """
    final_text = state.get('final_briefing')
    workflow_input = state.get('workflowInput', '未命名简报')  # 获取简报标题

    if not final_text:
        print("🚨 导出失败：final_briefing 内容为空。")
        return {"export_path": "导出失败：缺少最终内容"}

    # --- 1. 定义存储目录 ---
    DOCS_DIR = "generated_briefings"
    # 确保目录存在
    os.makedirs(DOCS_DIR, exist_ok=True)

    try:
        # --- 2. 生成唯一的、用户友好的文件名 ---
        # 清理标题，使其适用于文件名（去除文件系统不允许的特殊字符）
        safe_title = re.sub(r'[\\/:*?"<>|]', '', workflow_input).replace(' ', '_')
        # 添加时间戳以确保唯一性
        timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")

        # 最终文件名格式：[标题]_[时间戳].docx
        filename = f"{safe_title}_{timestamp}.docx"
        export_path = os.path.join(DOCS_DIR, filename)

        # --- 3. 文档生成逻辑 (基于你提供的代码) ---
        document = Document()

        # 简单处理：将文本按行分割，创建段落
        for line in final_text.split('\n'):
            line = line.strip()
            if not line:
                continue

            # 识别并设置标题/标题样式
            if line.startswith('##'):
                # 假设 ## 是二级标题
                document.add_heading(line.strip('#').strip(), level=2)
            elif line.startswith('#'):
                # 假设 # 是顶级标题
                document.add_heading(line.strip('#').strip(), level=1)
            else:
                # 普通段落
                document.add_paragraph(line)

        # 将文件保存到新的唯一路径
        document.save(export_path)

        print(f"✅ 简报已成功导出为 DOCX 文件: {export_path}")

        # 返回新的完整文件路径
        return {"export_path": export_path}

    except Exception as e:
        # 捕获并记录详细错误信息
        print(f"🚨 导出 DOCX 失败: {e}")
        return {"export_path": f"导出 DOCX 失败: {e}"}